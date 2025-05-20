import os
import filecmp
import difflib
import subprocess
import PyPDF2
import docx
import csv

def compare_files(file1_path, file2_path):
    """
    Compares two files and prints the differences.

    Args:
        file1_path (str): Path to the first file.
        file2_path (str): Path to the second file.
    """

    if not os.path.exists(file1_path) or not os.path.exists(file2_path):
        print("Error: One or both files do not exist.")
        return

    if filecmp.cmp(file1_path, file2_path, shallow=False):
        print("The files are identical.")
        return

    file1_extension = file1_path.split('.')[-1].lower()
    file2_extension = file2_path.split('.')[-1].lower()

    if file1_extension == file2_extension == "txt":
        compare_text_files(file1_path, file2_path)
    elif file1_extension == file2_extension == "pdf":
        compare_pdf_files(file1_path, file2_path)
    elif file1_extension == file2_extension == "docx":
        compare_docx_files(file1_path, file2_path)
    elif file1_extension == file2_extension == "csv":
        compare_csv_files(file1_path, file2_path)
    else:
        compare_binary_files(file1_path, file2_path)

def compare_text_files(file1_path, file2_path):
    """Compares two text files and prints the line differences."""

    with open(file1_path, 'r') as file1, open(file2_path, 'r') as file2:
        file1_lines = file1.readlines()
        file2_lines = file2.readlines()

    differ = difflib.Differ()
    diff = list(differ.compare(file1_lines, file2_lines))

    print("Differences found in text files:")
    for line in diff:
        if line.startswith('+ ') or line.startswith('- '):
            print(line)

def compare_pdf_files(file1_path, file2_path):
    """Compares two PDF files by extracting their text content."""

    try:
        with open(file1_path, 'rb') as file1, open(file2_path, 'rb') as file2:
            reader1 = PyPDF2.PdfReader(file1)
            reader2 = PyPDF2.PdfReader(file2)

            text1 = ""
            for page in reader1.pages:
                text1 += page.extract_text() + "\n"

            text2 = ""
            for page in reader2.pages:
                text2 += page.extract_text() + "\n"

        compare_text_files(text1.splitlines(), text2.splitlines())

    except Exception as e:
        print(f"Error comparing PDF files: {e}")
        print("Falling back to binary comparison.")
        compare_binary_files(file1_path, file2_path)

def compare_docx_files(file1_path, file2_path):
    """Compares two DOCX files by extracting their text content."""

    try:
        doc1 = docx.Document(file1_path)
        text1 = "\n".join([paragraph.text for paragraph in doc1.paragraphs])

        doc2 = docx.Document(file2_path)
        text2 = "\n".join([paragraph.text for paragraph in doc2.paragraphs])

        compare_text_files(text1.splitlines(), text2.splitlines())

    except Exception as e:
        print(f"Error comparing DOCX files: {e}")
        print("Falling back to binary comparison.")
        compare_binary_files(file1_path, file2_path)

def compare_csv_files(file1_path, file2_path):
    """Compares two CSV files and prints the differences."""

    try:
        with open(file1_path, 'r', newline='') as file1, open(file2_path, 'r', newline='') as file2:
            reader1 = csv.reader(file1)
            reader2 = csv.reader(file2)
            file1_lines = list(reader1)
            file2_lines = list(reader2)

        differ = difflib.Differ()
        diff = list(differ.compare(file1_lines, file2_lines))

        print("Differences found in CSV files:")
        for line in diff:
            if line.startswith('+ ') or line.startswith('- '):
                print(line)

    except Exception as e:
        print(f"Error comparing CSV files: {e}")
        print("Falling back to binary comparison.")
        compare_binary_files(file1_path, file2_path)

def compare_binary_files(file1_path, file2_path):
    """Compares two binary files using a basic byte-level comparison."""

    with open(file1_path, 'rb') as file1, open(file2_path, 'rb') as file2:
        file1_bytes = file1.read()
        file2_bytes = file2.read()

    if file1_bytes == file2_bytes:
        print("The files are identical (binary comparison).")
    else:
        print("The files are different (binary comparison).")

if __name__ == "__main__":
    file_path_1 = "visibility_outbound_transaction_report 18-05-2025.csv"  # Replace with your file paths
    file_path_2 = "visibility_outbound_transaction_report_2025-05-19.csv"
    compare_files(file_path_1, file_path_2)